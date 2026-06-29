"""
Serviços de proposta: render do template, geração de PDF (chrome) e DOCX (python-docx).

Fluxo:
  build_context(quotation)          -> dict de placeholders (empresa, cliente, feixe, preço)
  render_template_texts(tpl, ctx)   -> textos iniciais (Django Template, autoescape)
  create_proposal(quotation, tpl)   -> Proposal com textos editáveis pré-preenchidos
  generate(proposal)                -> escreve DOCX + PDF em MEDIA, grava hashes
"""
import hashlib
import os
import shutil
import subprocess
import tempfile
from datetime import date
from django.conf import settings
from django.core.files.storage import default_storage
from django.template import engines
from django.template.loader import render_to_string
from django.utils import timezone

from apps.proposals.models import Proposal, ProposalTemplate

_django_engine = engines["django"]


def build_context(quotation) -> dict:
    """Placeholders disponíveis nos templates de texto + dados da proposta."""
    inp = quotation.inputs or {}
    tpl = quotation.proposals.first().template if quotation.proposals.exists() else None
    base_tpl = tpl or ProposalTemplate.get_default()
    return {
        "empresa": base_tpl.company_name if base_tpl else "ENGEMATEX",
        "titulo": quotation.title,
        "escopo": quotation.get_scope_display(),
        "cliente": quotation.customer.company_name,
        "quotation_number": quotation.number,
        "n_tubos": inp.get("n_tubos", ""),
        "tubo_material": inp.get("tubo_material", ""),
        "tubo_od": inp.get("tubo_od_spec", ""),
        "tubo_parede": inp.get("tubo_wall_spec", ""),
        "tubo_comp_mm": inp.get("tubo_comp_mm", ""),
        "delivery_weeks": base_tpl.delivery_weeks if base_tpl else 8,
        "validity_days": base_tpl.validity_days if base_tpl else 30,
        "payment_terms": base_tpl.payment_terms if base_tpl else "30 dias",
        "preco_com_impostos": quotation.preco_com_impostos,
        "preco_sem_impostos": quotation.preco_sem_impostos,
    }


def _render(text: str, ctx: dict) -> str:
    # textos da proposta são TEXTO PLANO — sem autoescape na substituição de placeholders
    # (ex: OD 3/4" não vira 3/4&quot;). A escapagem p/ HTML acontece 1x ao exibir no template.
    wrapped = "{% autoescape off %}" + (text or "") + "{% endautoescape %}"
    return _django_engine.from_string(wrapped).render(ctx)


def render_template_texts(tpl: ProposalTemplate, ctx: dict) -> dict:
    return {
        "intro_text": _render(tpl.intro_template, ctx),
        "scope_text": _render(tpl.scope_template, ctx),
        "terms_text": _render(tpl.terms_template, ctx),
        "closing_text": _render(tpl.closing_template, ctx),
    }


def next_proposal_number(quotation) -> str:
    n = quotation.proposals.count() + 1
    letra = chr(ord("A") + n - 1)
    return f"PROP-{quotation.number.replace('COT-', '')}-{letra}"


def create_proposal(quotation, template: ProposalTemplate = None) -> Proposal:
    """Cria a proposta com os textos do template já renderizados (editáveis depois)."""
    template = template or ProposalTemplate.get_default()
    ctx = build_context(quotation)
    texts = render_template_texts(template, ctx) if template else {}
    return Proposal.objects.create(
        quotation=quotation, template=template,
        number=next_proposal_number(quotation), **texts)


# ---------------- Geração de arquivos ----------------

def _chrome_bin() -> str | None:
    for c in ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser"):
        p = shutil.which(c)
        if p:
            return p
    return None


def _pdf_context(proposal: Proposal) -> dict:
    """Contexto do template da proposta (inclui a memória de cálculo ASME como anexo)."""
    q = proposal.quotation
    return {"p": proposal, "q": q, "itens": q.itens.all(), "data": date.today(),
            "memorial": proposal_memorial(q)}


def _save_to_storage(local_path: str, storage_name: str) -> str:
    """Move um arquivo local para default_storage, retorna o storage name salvo."""
    if default_storage.exists(storage_name):
        default_storage.delete(storage_name)
    with open(local_path, "rb") as f:
        return default_storage.save(storage_name, f)


def _sha256_storage(name: str) -> str:
    h = hashlib.sha256()
    with default_storage.open(name, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def generate_pdf(proposal: Proposal) -> str:
    """Renderiza a proposta como HTML (design G), converte para PDF e salva via default_storage.
    Usa WeasyPrint se disponível, senão chrome headless --print-to-pdf."""
    html = render_to_string("proposals/proposal_pdf.html", _pdf_context(proposal))
    storage_name = f"proposals/{proposal.number}.pdf"
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
    os.close(tmp_fd)
    try:
        written = False
        try:
            from weasyprint import HTML  # lazy (libs de sistema)
            HTML(string=html).write_pdf(tmp_path)
            written = True
        except Exception:
            pass
        if not written:
            chrome = _chrome_bin()
            if not chrome:
                raise RuntimeError("Sem backend de PDF (weasyprint ou chrome).")
            with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False, encoding="utf-8") as fp:
                fp.write(html)
                tmp_html = fp.name
            try:
                subprocess.run([chrome, "--headless=new", "--no-sandbox", "--disable-gpu",
                                "--no-pdf-header-footer", f"--print-to-pdf={tmp_path}",
                                f"file://{tmp_html}"],
                               check=True, capture_output=True, timeout=60)
            finally:
                os.unlink(tmp_html)
        return _save_to_storage(tmp_path, storage_name)
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def generate_docx(proposal: Proposal) -> str:
    """Gera o DOCX (editável no Word) e salva via default_storage. Retorna o storage name."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    q = proposal.quotation
    doc = Document()
    title = doc.add_heading(level=0)
    run = title.add_run(f"PROPOSTA COMERCIAL — {proposal.number}")
    run.font.color.rgb = RGBColor(0x16, 0x15, 0x1A)

    meta = doc.add_paragraph()
    meta.add_run(f"Cliente: {q.customer.company_name}\n").bold = True
    meta.add_run(f"Cotação: {q.number}   Data: {date.today():%d/%m/%Y}")

    for label, text in (("", proposal.intro_text), ("ESCOPO", proposal.scope_text)):
        if label:
            doc.add_heading(label, level=1)
        doc.add_paragraph(text)

    doc.add_heading("COMPOSIÇÃO", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Código", "Item", "Material (R$)", "Mão de Obra (R$)")):
        hdr[i].text = h
    for it in q.itens.all():
        row = table.add_row().cells
        row[0].text = it.codigo_item
        row[1].text = it.descricao
        row[2].text = f"{it.custo_material:,.2f}"
        row[3].text = f"{it.custo_mo:,.2f}"

    price = doc.add_paragraph()
    price.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr = price.add_run(f"PREÇO DE VENDA: R$ {q.preco_com_impostos:,.2f}")
    pr.bold = True
    pr.font.size = Pt(14)

    memorial = proposal_memorial(q)
    if memorial:
        doc.add_heading("ANEXO — Memória de cálculo ASME", level=1)
        mt = doc.add_table(rows=1, cols=3)
        mt.style = "Light Grid Accent 1"
        mh = mt.rows[0].cells
        for i, h in enumerate(("Verificação", "Resultado", "Norma / Fonte")):
            mh[i].text = h
        for e in memorial:
            row = mt.add_row().cells
            row[0].text = e["item"] + (f" [{e['status']}]" if e.get("status") else "")
            row[1].text = e.get("valor", "")
            row[2].text = e.get("norma", "") + (f" — {e['fonte']}" if e.get("fonte") else "")
        doc.add_paragraph("Anexo de apoio à certificação — não substitui o memorial do "
                          "engenheiro responsável. Itens de estimativa estão marcados.")

    doc.add_heading("CONDIÇÕES", level=1)
    doc.add_paragraph(proposal.terms_text)
    doc.add_paragraph(proposal.closing_text)

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    try:
        doc.save(tmp_path)
        return _save_to_storage(tmp_path, f"proposals/{proposal.number}.docx")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def generate(proposal: Proposal) -> Proposal:
    """Gera DOCX + PDF via default_storage, grava storage names + hashes, marca como pronta."""
    docx_name = generate_docx(proposal)
    pdf_name = generate_pdf(proposal)
    proposal.docx_path = docx_name
    proposal.pdf_path = pdf_name
    proposal.docx_sha256 = _sha256_storage(docx_name)
    proposal.pdf_sha256 = _sha256_storage(pdf_name)
    proposal.status = "ready"
    proposal.generated_at = timezone.now()
    proposal.save()
    return proposal


def proposal_memorial(quotation):
    """Memória de cálculo ASME da cotação (só p/ equipamento completo/permutador). [] p/ feixe.
    Anexo de certificação que acompanha a proposta — reusa tema_templates.memorial_asme."""
    if getattr(quotation, "scope", None) != "complete":
        return []
    inp = quotation.inputs or {}
    from apps.tema_templates.services import memorial_asme
    return memorial_asme(inp.get("designacao", ""), inp)
