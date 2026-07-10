"""
Testes do app proposals: template configurável, customização por caso, geração DOCX/PDF.
"""
import os
import shutil
import unittest
from unittest.mock import patch
from django.contrib.auth.models import User
from django.core import mail
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.template import engines
from django.test import SimpleTestCase, override_settings
from django_tenants.test.cases import TenantTestCase

from apps.quotations.models import Customer
from apps.quotations.services import create_feixe_quotation
from apps.proposals.models import Proposal, ProposalTemplate, ProposalVersion
from apps.proposals import services
from apps.audit.models import AccessLog


def _chrome_ou_weasy() -> bool:
    if any(shutil.which(c) for c in ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser")):
        return True
    try:
        import weasyprint  # noqa
        return True
    except Exception:
        return False


class ProposalTests(TenantTestCase):
    def setUp(self):
        self.tpl = ProposalTemplate.objects.create(name="Padrão", is_default=True)
        self.customer = Customer.objects.create(company_name="Petrobras RPBC")
        self.q = create_feixe_quotation(self.customer, "Feixe 136 tubos")

    def test_template_renderiza_placeholders(self):
        ctx = services.build_context(self.q)
        texts = services.render_template_texts(self.tpl, ctx)
        self.assertIn("Feixe 136 tubos", texts["intro_text"])     # {{ titulo }}
        self.assertIn("136", texts["scope_text"])                  # {{ n_tubos }}
        self.assertIn("8 semanas", texts["terms_text"])            # {{ delivery_weeks }}
        # texto plano: OD 3/4" não pode virar 3/4&quot; (sem autoescape na substituição)
        self.assertIn('3/4"', texts["scope_text"])
        self.assertNotIn("&quot;", texts["scope_text"])

    def test_create_proposal_pre_preenche_textos_editaveis(self):
        p = services.create_proposal(self.q, self.tpl)
        self.assertTrue(p.intro_text)                              # veio do template
        self.assertTrue(p.number.startswith("PROP-"))
        self.assertEqual(p.status, "draft")

    def test_customizacao_por_caso_nao_altera_template(self):
        p = services.create_proposal(self.q, self.tpl)
        p.intro_text = "Texto customizado para ESTE cliente."
        p.save()
        # o template-modelo permanece intacto
        self.tpl.refresh_from_db()
        self.assertIn("{{ titulo }}", self.tpl.intro_template)
        self.assertEqual(Proposal.objects.get(pk=p.pk).intro_text, "Texto customizado para ESTE cliente.")

    def test_generate_docx_usa_default_storage(self):
        """generate_docx() deve usar default_storage.save, não escrever direto no filesystem."""
        p = services.create_proposal(self.q, self.tpl)
        with patch.object(default_storage, 'save', wraps=default_storage.save) as mock_save:
            name = services.generate_docx(p)
        mock_save.assert_called()
        # retorna nome de storage relativo (não caminho absoluto)
        self.assertFalse(os.path.isabs(name))
        self.assertTrue(default_storage.exists(name))
        default_storage.delete(name)

    def test_generate_docx_prefixa_por_tenant_schema(self):
        """O storage name deve conter o schema do tenant → isolamento no MEDIA/bucket
        compartilhado. Sem o prefixo, dois tenants com o mesmo proposal.number
        colidiriam e um baixaria o arquivo do outro."""
        from django.db import connection
        p = services.create_proposal(self.q, self.tpl)
        name = services.generate_docx(p)
        self.assertIn(f"proposals/{connection.schema_name}/", name)
        self.assertTrue(name.endswith(f"/{p.number}.docx"))
        default_storage.delete(name)

    def test_generate_docx_cria_arquivo_e_hash(self):
        p = services.create_proposal(self.q, self.tpl)
        name = services.generate_docx(p)
        self.assertTrue(default_storage.exists(name))
        with default_storage.open(name, "rb") as f:
            self.assertGreater(len(f.read()), 1000)
        default_storage.delete(name)

    @unittest.skipUnless(_chrome_ou_weasy(), "sem backend de PDF (chrome/weasyprint)")
    def test_generate_completo_docx_pdf_hashes(self):
        p = services.create_proposal(self.q, self.tpl)
        services.generate(p)
        self.assertEqual(p.status, "ready")
        self.assertTrue(p.pdf_path and p.docx_path)
        self.assertEqual(len(p.pdf_sha256), 64)
        self.assertEqual(len(p.docx_sha256), 64)

    def test_generate_cria_entrada_no_historico(self):
        p = services.create_proposal(self.q, self.tpl)
        with patch.object(services, "generate_docx", return_value="proposals/test/proposta.docx"), \
                patch.object(services, "generate_pdf", return_value="proposals/test/proposta.pdf"), \
                patch.object(services, "_sha256_storage", side_effect=["d" * 64, "p" * 64]):
            services.generate(p)
        versao = ProposalVersion.objects.get(proposal=p)
        self.assertEqual(versao.version_number, 1)
        self.assertEqual(versao.number, p.number)
        self.assertEqual(versao.pdf_path, "proposals/test/proposta.pdf")
        self.assertIsNotNone(versao.generated_at)

    def test_numeracao_proposta_por_revisao(self):
        p1 = services.create_proposal(self.q, self.tpl)
        p2 = services.create_proposal(self.q, self.tpl)
        self.assertTrue(p1.number.endswith("-A"))
        self.assertTrue(p2.number.endswith("-B"))

    def test_objeto_exclusoes_herdados_do_template(self):
        """§4 Objeto e §5 Exclusões: proposta criada do template herda os textos (snapshot)."""
        p = services.create_proposal(self.q, self.tpl)
        self.assertTrue(p.object_text)        # veio do object_template
        self.assertTrue(p.exclusions_text)    # veio do exclusions_template

    def test_pdf_context_inclui_objeto_e_exclusoes(self):
        from django.template.loader import render_to_string
        p = services.create_proposal(self.q, self.tpl)
        p.object_text = "OBJETO-UNICO-XYZ"
        p.exclusions_text = "EXCLUSAO-UNICA-QWE"
        p.save()
        html = render_to_string("proposals/proposal_pdf.html", services._pdf_context(p))
        self.assertIn("Objeto da Proposta", html)
        self.assertIn("OBJETO-UNICO-XYZ", html)
        self.assertIn("Exclusões do Fornecimento", html)
        self.assertIn("EXCLUSAO-UNICA-QWE", html)

    def test_pdf_preserva_formatacao_rich_text_sanitizada_em_objeto_e_exclusoes(self):
        from django.template.loader import render_to_string
        p = services.create_proposal(self.q, self.tpl)
        p.object_text = "<p>Escopo <strong>principal</strong></p><script>alert(1)</script>"
        p.exclusions_text = "<ul><li>Frete</li><li><em>Montagem</em></li></ul>"
        p.save()

        html = render_to_string("proposals/proposal_pdf.html", services._pdf_context(p))

        self.assertIn("<strong>principal</strong>", html)
        self.assertIn("<ul><li>Frete</li><li><em>Montagem</em></li></ul>", html)
        self.assertNotIn("<script>", html)
        self.assertNotIn("alert(1)", html)


class ProposalRichTextUnitTests(SimpleTestCase):
    def test_proposal_richtext_filter_preserva_html_permitido_e_remove_script(self):
        tpl = engines["django"].from_string(
            "{% load proposal_richtext %}{{ value|proposal_richtext }}"
        )

        html = tpl.render({"value": "<p>Escopo <strong>ok</strong></p><script>alert(1)</script>"})

        self.assertIn("<strong>ok</strong>", html)
        self.assertNotIn("<script>", html)
        self.assertNotIn("alert(1)", html)

    def test_edit_template_contem_controles_rich_text_para_objeto_e_exclusoes(self):
        with open("apps/proposals/templates/proposals/edit.html", encoding="utf-8") as fh:
            html = fh.read()

        self.assertEqual(html.count('contenteditable="true"'), 2)
        self.assertIn('data-rich-text="object_text"', html)
        self.assertIn('data-rich-text="exclusions_text"', html)
        self.assertGreaterEqual(html.count('data-rich-action="bold"'), 2)


class ProposalViewTests(TenantTestCase):
    def setUp(self):
        self.client.defaults["HTTP_HOST"] = self.get_test_tenant_domain()
        self.user = User.objects.create_user(username="orc", password="senha-forte-123")
        from apps.accounts.models import UserProfile
        # Criar/editar/gerar proposta exige papel de escrita (RBAC H2.8: Engenheiro/Admin).
        UserProfile.objects.create(user=self.user, full_name="Orc", role=UserProfile.ROLE_ADMIN)
        self.client.force_login(self.user)
        ProposalTemplate.objects.create(name="Padrão", is_default=True)
        self.q = create_feixe_quotation(Customer.objects.create(company_name="Cli"), "Feixe")

    def test_fluxo_criar_editar(self):
        # criar a partir do template -> redireciona pro editor
        resp = self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        self.assertEqual(resp.status_code, 302)
        p = Proposal.objects.latest("created_at")
        # editar customizando o texto
        resp = self.client.post(f"/propostas/{p.pk}/editar/", {
            "intro_text": "Intro customizada", "scope_text": "Escopo",
            "terms_text": "Condições", "closing_text": "Fecho", "save": "1"})
        self.assertEqual(resp.status_code, 302)
        p.refresh_from_db()
        self.assertEqual(p.intro_text, "Intro customizada")

    def test_edicao_persiste_objeto_e_exclusoes(self):
        self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        p = Proposal.objects.latest("created_at")
        resp = self.client.post(f"/propostas/{p.pk}/editar/", {
            "intro_text": "i", "scope_text": "s", "terms_text": "t", "closing_text": "c",
            "object_text": "Objeto customizado", "exclusions_text": "Exclusão customizada",
            "save": "1"})
        self.assertEqual(resp.status_code, 302)
        p.refresh_from_db()
        self.assertEqual(p.object_text, "Objeto customizado")
        self.assertEqual(p.exclusions_text, "Exclusão customizada")


    def test_edit_expoe_templates_no_contexto(self):
        """Tela 06: o editor expõe os textos-modelo (§4 Objeto / §5 Exclusões) do
        ProposalTemplate no contexto, para o botão 'Carregar modelo' injetar via Alpine."""
        self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        p = Proposal.objects.latest("created_at")
        resp = self.client.get(f"/propostas/{p.pk}/editar/")
        self.assertEqual(resp.status_code, 200)
        self.assertIn("object_template", resp.context)
        self.assertIn("exclusions_template", resp.context)
        # a cláusula-modelo de exclusões traz o boilerplate de frete/montagem
        self.assertIn("frete", resp.context["exclusions_template"])
        self.assertIn("Fornecimento de", resp.context["object_template"])
        # e o controle de carga aparece no HTML renderizado
        self.assertContains(resp, "Carregar modelo")

    def test_edit_renderiza_controles_rich_text_para_objeto_e_exclusoes(self):
        self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        p = Proposal.objects.latest("created_at")

        resp = self.client.get(f"/propostas/{p.pk}/editar/")

        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'contenteditable="true"', count=2)
        self.assertContains(resp, 'data-rich-text="object_text"')
        self.assertContains(resp, 'data-rich-text="exclusions_text"')
        self.assertContains(resp, 'data-rich-action="bold"', count=2)

    def test_generate_view_registra_access_log(self):
        p = services.create_proposal(self.q)
        resp = self.client.post(f"/propostas/{p.pk}/editar/", {
            "intro_text": "Intro", "scope_text": "Escopo",
            "terms_text": "Condições", "closing_text": "Fecho", "generate": "1"})
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(AccessLog.objects.filter(action="generate", resource_id=str(p.pk)).exists())

    def test_download_usa_default_storage(self):
        """proposal_download deve abrir o arquivo via default_storage, não os.path direto."""
        p = services.create_proposal(self.q)
        storage_name = "proposals/test-ds-storage.docx"
        p.docx_path = storage_name
        p.save()
        if default_storage.exists(storage_name):
            default_storage.delete(storage_name)
        default_storage.save(storage_name, ContentFile(b"docx-test"))
        try:
            with patch.object(default_storage, 'open', wraps=default_storage.open) as mock_open:
                resp = self.client.get(f"/propostas/{p.pk}/download/docx/")
            self.assertEqual(resp.status_code, 200)
            mock_open.assert_called_once_with(storage_name, "rb")
        finally:
            if default_storage.exists(storage_name):
                default_storage.delete(storage_name)

    def test_download_registra_access_log(self):
        p = services.create_proposal(self.q)
        storage_name = "proposals/test-download.docx"
        p.docx_path = storage_name
        p.save()
        if default_storage.exists(storage_name):
            default_storage.delete(storage_name)
        default_storage.save(storage_name, ContentFile(b"docx"))
        try:
            resp = self.client.get(f"/propostas/{p.pk}/download/docx/")
            self.assertEqual(resp.status_code, 200)
            self.assertTrue(AccessLog.objects.filter(action="download", resource_id=str(p.pk)).exists())
        finally:
            if default_storage.exists(storage_name):
                default_storage.delete(storage_name)


    def test_download_formato_invalido_retorna_404(self):
        p = services.create_proposal(self.q)
        resp = self.client.get(f"/propostas/{p.pk}/download/xlsx/")
        self.assertEqual(resp.status_code, 404)

    @override_settings(
        EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend",
        DEFAULT_FROM_EMAIL="orcamentos@example.com",
    )
    def test_send_email_envia_pdf_anexado_e_marca_sent(self):
        p = services.create_proposal(self.q)
        storage_name = "proposals/test/send-email.pdf"
        p.pdf_path = storage_name
        p.save(update_fields=["pdf_path"])
        default_storage.save(storage_name, ContentFile(b"%PDF-test"))
        ProposalVersion.objects.create(proposal=p, version_number=1, number=p.number, pdf_path=storage_name)
        try:
            resp = self.client.post(f"/propostas/{p.pk}/enviar-email/", {
                "to_email": "cliente@example.com",
                "body": "Segue proposta anexa.",
            }, follow=True)
            self.assertEqual(resp.status_code, 200)
            p.refresh_from_db()
            self.assertEqual(p.status, "sent")
            self.assertEqual(len(mail.outbox), 1)
            message = mail.outbox[0]
            self.assertEqual(message.to, ["cliente@example.com"])
            self.assertEqual(message.from_email, "orcamentos@example.com")
            self.assertEqual(message.body, "Segue proposta anexa.")
            self.assertEqual(len(message.attachments), 1)
            attachment = message.attachments[0]
            self.assertEqual(attachment[0], f"{p.number}.pdf")
            self.assertEqual(attachment[2], "application/pdf")
            versao = ProposalVersion.objects.get(proposal=p)
            self.assertIsNotNone(versao.emailed_at)
            self.assertEqual(versao.email_to, "cliente@example.com")
            self.assertContains(resp, "E-mail enviado com sucesso.")
        finally:
            if default_storage.exists(storage_name):
                default_storage.delete(storage_name)

    def test_send_email_destinatario_invalido_falha_com_mensagem(self):
        p = services.create_proposal(self.q)
        resp = self.client.post(f"/propostas/{p.pk}/enviar-email/", {
            "to_email": "destinatario-invalido",
            "body": "Segue proposta anexa.",
        }, follow=True)
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, "Informe um e-mail de destino válido.")
        p.refresh_from_db()
        self.assertNotEqual(p.status, "sent")

    @override_settings(EMAIL_BACKEND="django.core.mail.backends.smtp.EmailBackend")
    def test_send_email_sem_backend_funcional_falha_graciosamente(self):
        p = services.create_proposal(self.q)
        storage_name = "proposals/test/send-email-error.pdf"
        p.pdf_path = storage_name
        p.save(update_fields=["pdf_path"])
        default_storage.save(storage_name, ContentFile(b"%PDF-test"))
        try:
            with patch("apps.proposals.services.EmailMessage.send", side_effect=OSError("smtp down")):
                resp = self.client.post(f"/propostas/{p.pk}/enviar-email/", {
                    "to_email": "cliente@example.com",
                    "body": "Segue proposta anexa.",
                }, follow=True)
            self.assertEqual(resp.status_code, 200)
            self.assertContains(resp, "Nao foi possivel enviar o e-mail com a configuracao atual.")
            p.refresh_from_db()
            self.assertNotEqual(p.status, "sent")
        finally:
            if default_storage.exists(storage_name):
                default_storage.delete(storage_name)

    def test_detail_e_edit_exibem_historico_e_botao_de_email(self):
        p = services.create_proposal(self.q)
        ProposalVersion.objects.create(
            proposal=p, version_number=1, number=p.number, pdf_path="proposals/test/v1.pdf"
        )
        detail = self.client.get(f"/propostas/{p.pk}/")
        edit = self.client.get(f"/propostas/{p.pk}/editar/")
        self.assertContains(detail, "Historico de versoes")
        self.assertContains(detail, "Enviar por E-mail")
        self.assertContains(edit, "Historico de versoes")
        self.assertContains(edit, "Enviar por E-mail")

    def test_login_required(self):
        self.client.logout()
        resp = self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        self.assertEqual(resp.status_code, 302)
        self.assertIn("/login/", resp.url)

    def test_editor_mostra_botao_gerar_e_visualizar(self):
        """Tela 06 item 3: o botão final de emissão foi trocado para 'GERAR E VISUALIZAR PDF'."""
        self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        p = Proposal.objects.latest("created_at")
        resp = self.client.get(f"/propostas/{p.pk}/editar/")
        self.assertContains(resp, "GERAR E VISUALIZAR PDF")

    def test_preview_gera_pdf_e_serve_inline_quando_ainda_nao_existe(self):
        """Rota de preview: sem exigir bytes reais de PDF (mock de generate_pdf) — valida
        que a rota gera (se necessário) e serve com Content-Disposition inline."""
        p = services.create_proposal(self.q)
        self.assertFalse(p.pdf_path)
        fake_name = "proposals/test/preview-fake.pdf"
        default_storage.save(fake_name, ContentFile(b"%PDF-fake-preview"))
        try:
            with patch.object(services, "generate_pdf", return_value=fake_name) as mock_gen:
                resp = self.client.get(f"/propostas/{p.pk}/preview/")
            mock_gen.assert_called_once()
            self.assertEqual(resp.status_code, 200)
            self.assertIn("inline", resp["Content-Disposition"])
            p.refresh_from_db()
            self.assertEqual(p.pdf_path, fake_name)
        finally:
            if default_storage.exists(fake_name):
                default_storage.delete(fake_name)

    def test_preview_reusa_pdf_existente_sem_regenerar(self):
        p = services.create_proposal(self.q)
        storage_name = "proposals/test/preview-existing.pdf"
        default_storage.save(storage_name, ContentFile(b"%PDF-existing"))
        p.pdf_path = storage_name
        p.save()
        try:
            with patch.object(services, "generate_pdf") as mock_gen:
                resp = self.client.get(f"/propostas/{p.pk}/preview/")
            mock_gen.assert_not_called()
            self.assertEqual(resp.status_code, 200)
            self.assertIn("inline", resp["Content-Disposition"])
        finally:
            default_storage.delete(storage_name)

    def test_preview_registra_access_log(self):
        p = services.create_proposal(self.q)
        storage_name = "proposals/test/preview-log.pdf"
        default_storage.save(storage_name, ContentFile(b"%PDF"))
        p.pdf_path = storage_name
        p.save()
        try:
            self.client.get(f"/propostas/{p.pk}/preview/")
            self.assertTrue(AccessLog.objects.filter(action="preview", resource_id=str(p.pk)).exists())
        finally:
            default_storage.delete(storage_name)

    @unittest.skipUnless(_chrome_ou_weasy(), "sem backend de PDF (chrome/weasyprint)")
    def test_fluxo_completo_gerar_mostra_preview_iframe(self):
        """Fluxo real: gerar (chrome/weasyprint) -> redireciona -> página mostra iframe de preview
        apontando pra rota inline, e essa rota responde servindo o PDF inline."""
        self.client.get(f"/cotacoes/{self.q.pk}/proposta/nova/")
        p = Proposal.objects.latest("created_at")
        resp = self.client.post(f"/propostas/{p.pk}/editar/", {
            "intro_text": "Intro", "scope_text": "Escopo",
            "terms_text": "Condições", "closing_text": "Fecho", "generate": "1"}, follow=True)
        self.assertEqual(resp.status_code, 200)
        p.refresh_from_db()
        self.assertTrue(p.pdf_path)
        preview_url = f"/propostas/{p.pk}/preview/"
        self.assertContains(resp, preview_url)
        self.assertContains(resp, "<iframe")
        preview_resp = self.client.get(preview_url)
        self.assertEqual(preview_resp.status_code, 200)
        self.assertIn("inline", preview_resp["Content-Disposition"])


class ProposalMemorialTests(TenantTestCase):
    """Memorial ASME embutido na proposta (anexo de certificação no PDF)."""

    def _permutador_q(self):
        from apps.quotations.services import create_permutador_quotation
        from pricing_engine.permutador_quote import quote_completo
        cust = Customer.objects.create(company_name="ACME")
        cleaned = {"classe_casco": "CS", "pressao_projeto_bar": 50, "temperatura_projeto_c": 150,
                   "rt_escopo": "Total", "diametro_casco_mm": 764, "esp_casco_mm": 9.5,
                   "corrosao_mm": 3, "comprimento_casco_mm": 1631}
        return create_permutador_quotation(cust, "BEU", cleaned, quote_completo("BEU"))

    def test_proposal_memorial_para_permutador(self):
        from apps.proposals.services import proposal_memorial
        memo = proposal_memorial(self._permutador_q())
        blob = " ".join(e["item"] for e in memo)
        self.assertIn("UG-27", blob)
        self.assertTrue(any("2025" in (e.get("fonte") or "") for e in memo))

    def test_proposal_memorial_feixe_vazio(self):
        from apps.proposals.services import proposal_memorial
        from apps.quotations.services import create_feixe_quotation
        cust = Customer.objects.create(company_name="X Ltda")
        q = create_feixe_quotation(cust, "Feixe")
        self.assertEqual(proposal_memorial(q), [])

    def test_proposal_pdf_html_inclui_memorial(self):
        from apps.proposals.services import create_proposal, _pdf_context
        from django.template.loader import render_to_string
        prop = create_proposal(self._permutador_q())
        html = render_to_string("proposals/proposal_pdf.html", _pdf_context(prop))
        self.assertIn("Memória de cálculo", html)
        self.assertIn("UG-27", html)
        self.assertIn("2025", html)            # procedência normativa no anexo

    def test_proposal_docx_inclui_memorial(self):
        from apps.proposals.services import create_proposal, generate_docx
        import docx
        prop = create_proposal(self._permutador_q())
        name = generate_docx(prop)
        try:
            with default_storage.open(name, "rb") as f:
                d = docx.Document(f)
            full = "\n".join(p.text for p in d.paragraphs)
            for t in d.tables:
                for row in t.rows:
                    full += "\n" + " ".join(c.text for c in row.cells)
            self.assertIn("Memória de cálculo", full)
            self.assertIn("UG-27", full)
        finally:
            if default_storage.exists(name):
                default_storage.delete(name)
